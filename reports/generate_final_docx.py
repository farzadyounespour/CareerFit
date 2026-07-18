from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from generate_midterm_docx import (
    INK,
    SLATE,
    TEAL,
    add_architecture_diagram,
    add_data_model_diagram,
    add_markdown_content,
    add_navigation_diagram,
    add_toc,
    add_wireframes,
    add_workflow_diagram,
    configure_document,
    set_cell_text,
    style_table,
)


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "final-report.md"
OUTPUT_PATH = BASE_DIR / "CareerFit_Final_Report_Draft.docx"


def add_cover_page(document):
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CareerFit")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Comparative and Explainable Resume-Job Matching Framework")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(INK)

    report_type = document.add_paragraph()
    report_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_type.add_run("Final Project Report Draft")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(SLATE)

    document.add_paragraph()
    document.add_paragraph()
    details = [
        ("Student", "Farzad Younespour"),
        ("Student ID", "40306504"),
        ("Course", "Project and Report I"),
        ("Supervisor", "Professor Joumana Dargham"),
        ("Term", "Summer 2026"),
    ]
    table = document.add_table(rows=len(details), cols=2)
    style_table(table)
    for index, (label, value) in enumerate(details):
        set_cell_text(table.cell(index, 0), label, bold=True, color=TEAL, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(index, 1), value, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Working draft: refresh final screenshots, diagrams, and case-study discussion before submission.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    document.add_page_break()


def main():
    markdown = SOURCE_PATH.read_text(encoding="utf-8")
    abstract_start = markdown.index("## Abstract")
    toc_start = markdown.index("## Table of Contents", abstract_start)
    introduction_start = markdown.index("## 1. Introduction", toc_start)
    markdown = markdown[abstract_start:toc_start] + markdown[introduction_start:]
    document = Document()
    configure_document(document)
    add_cover_page(document)
    document.add_heading("Table of Contents", level=1)
    add_toc(document.add_paragraph())
    document.add_page_break()
    add_markdown_content(document, markdown)
    document.add_page_break()
    add_architecture_diagram(document, appendix="F")
    add_workflow_diagram(document, appendix="G")
    add_navigation_diagram(document, appendix="H")
    add_data_model_diagram(document, appendix="I")
    add_wireframes(document, appendix="J")
    document.save(OUTPUT_PATH)
    print(f"Generated {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
