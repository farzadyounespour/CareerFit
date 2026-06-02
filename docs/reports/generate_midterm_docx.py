import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "midterm-report.md"
OUTPUT_PATH = BASE_DIR / "CareerFit_Midterm_Report.docx"

TEAL = "0F766E"
INK = "172033"
SLATE = "64748B"
LIGHT_TEAL = "ECFDF5"
LIGHT_BLUE = "EFF6FF"
LIGHT_SLATE = "F8FAFC"
LIGHT_AMBER = "FFFBEB"
LIGHT_ROSE = "FFF1F2"
LIGHT_PURPLE = "F5F3FF"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def add_toc(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update this field in Word to generate the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([field, instruction, separate, placeholder, end])


def clean_inline(text):
    return re.sub(r"[*_`]", "", text).strip()


def add_markdown_table(document, lines):
    rows = []
    for line in lines:
        values = [clean_inline(value.strip()) for value in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", value.replace(" ", "")) for value in values):
            continue
        rows.append(values)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    style_table(table)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            set_cell_text(
                cell,
                value,
                bold=row_index == 0,
                color=INK,
                size=8 if row_index else 9,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            if row_index == 0:
                set_cell_shading(cell, LIGHT_TEAL)
    document.add_paragraph()


def add_markdown_content(document, markdown):
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = "No Spacing"
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if stripped == "---" or not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:]), level=2)
        elif re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        elif stripped.startswith("- "):
            document.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph(clean_inline(stripped[2:]))
            paragraph.style = "Intense Quote"
        elif stripped.startswith("**") and stripped.endswith("**"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(clean_inline(stripped))
            run.bold = True
        else:
            document.add_paragraph(clean_inline(stripped))
        index += 1


def add_cover_page(document):
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CareerFit")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Comparative and Explainable Resume-Job Matching Framework")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(INK)

    report_type = document.add_paragraph()
    report_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_type.add_run("Midterm Project Report")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(SLATE)

    document.add_paragraph()
    document.add_paragraph()
    details = [
        ("Student", "Farzad Younespour"),
        ("Student ID", "40306504"),
        ("Course", "Project and Report I"),
        ("Supervisor", "Professor Joumana Dargham"),
        ("Term", "Summer 2026"),
    ]
    table = document.add_table(rows=len(details), cols=2)
    style_table(table)
    for index, (label, value) in enumerate(details):
        set_cell_text(table.cell(index, 0), label, bold=True, color=TEAL, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(index, 1), value, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Design Phase: Proposed Features, Wireframes, Architecture, and Implementation Plan")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    document.add_page_break()


def add_architecture_diagram(document, appendix="A"):
    document.add_heading(f"Appendix {appendix}: Editable Architecture Diagram", level=1)
    document.add_paragraph(f"Figure {appendix}1. Proposed CareerFit client-server architecture. The diagram is editable in Word.")
    table = document.add_table(rows=7, cols=3)
    style_table(table)
    cells = [row.cells for row in table.rows]

    cells[0][0].merge(cells[0][2])
    set_cell_text(cells[0][0], "React Web Frontend\nHome | Profile | Resume | Jobs | Report | Tracker", bold=True, size=10)
    set_cell_shading(cells[0][0], LIGHT_BLUE)
    cells[1][0].merge(cells[1][2])
    set_cell_text(cells[1][0], "REST API requests and responses", bold=True, color=TEAL)
    cells[2][0].merge(cells[2][2])
    set_cell_text(cells[2][0], "Django REST Framework Backend", bold=True, size=10)
    set_cell_shading(cells[2][0], LIGHT_TEAL)
    for index, label in enumerate(("Accounts", "Resumes", "Jobs and Tracker")):
        set_cell_text(cells[3][index], label, bold=True, size=9)
        set_cell_shading(cells[3][index], LIGHT_SLATE)
    cells[4][0].merge(cells[4][2])
    set_cell_text(cells[4][0], "Matching, ATS Checks, Recommendations, Optional AI Coaching", bold=True, size=9)
    set_cell_shading(cells[4][0], LIGHT_AMBER)
    cells[5][0].merge(cells[5][2])
    set_cell_text(cells[5][0], "SQLite locally / PostgreSQL for deployment", bold=True, size=9)
    set_cell_shading(cells[5][0], LIGHT_BLUE)
    for index, label in enumerate(("Adzuna API", "Arbeitnow API", "Optional Jooble / Ollama / OpenAI")):
        set_cell_text(cells[6][index], label, bold=True, size=8)
        set_cell_shading(cells[6][index], LIGHT_SLATE)
    document.add_paragraph()


def add_workflow_diagram(document, appendix="B"):
    document.add_heading(f"Appendix {appendix}: Editable User Workflow", level=1)
    document.add_paragraph(f"Figure {appendix}1. Proposed CareerFit workflow. Each box is editable in Word.")
    steps = [
        "1. Create account or sign in",
        "2. Complete candidate profile",
        "3. Upload or paste resume",
        "4. Review editable resume text and ATS checks",
        "5. Search jobs, import URL, or paste description",
        "6. Select a posting",
        "7. Review quick comparison",
        "8. Generate detailed readiness report",
        "9. Improve resume and rescan OR add role to tracker",
    ]
    table = document.add_table(rows=(len(steps) * 2) - 1, cols=1)
    style_table(table)
    for index, step in enumerate(steps):
        row_index = index * 2
        set_cell_text(table.cell(row_index, 0), step, bold=True, size=9)
        set_cell_shading(table.cell(row_index, 0), LIGHT_TEAL if index % 2 == 0 else LIGHT_BLUE)
        if row_index + 1 < len(table.rows):
            set_cell_text(table.cell(row_index + 1, 0), "v", bold=True, color=TEAL, size=10)
    document.add_paragraph()


def add_navigation_diagram(document, appendix="C"):
    document.add_heading(f"Appendix {appendix}: Editable Navigation Map", level=1)
    document.add_paragraph(f"Figure {appendix}1. Proposed website information architecture. Each item is editable in Word.")
    table = document.add_table(rows=5, cols=6)
    style_table(table)
    cells = [row.cells for row in table.rows]

    cells[0][0].merge(cells[0][5])
    set_cell_text(cells[0][0], "CareerFit Web Application", bold=True, size=11)
    set_cell_shading(cells[0][0], LIGHT_TEAL)
    cells[1][0].merge(cells[1][5])
    set_cell_text(cells[1][0], "Horizontal navigation", bold=True, color=TEAL, size=9)
    pages = ("Home", "Profile", "Resume", "Jobs", "Report", "Tracker")
    descriptions = (
        "Overview\nStarting actions",
        "Preferences\nAutofill data",
        "Upload\nEdit\nSave versions",
        "Search\nImport\nQuick compare",
        "Scores\nEvidence\nImprove",
        "Stages\nTasks\nFollow-ups",
    )
    for column, page in enumerate(pages):
        set_cell_text(cells[2][column], page, bold=True, size=9)
        set_cell_shading(cells[2][column], LIGHT_BLUE if column % 2 == 0 else LIGHT_SLATE)
        set_cell_text(cells[3][column], descriptions[column], size=8)
    cells[4][0].merge(cells[4][5])
    set_cell_text(
        cells[4][0],
        "Public entry points: Home and authentication | Private workspace records: profile, resume versions, reports, and tracked jobs",
        size=8,
    )
    set_cell_shading(cells[4][0], LIGHT_AMBER)
    document.add_paragraph()


def add_data_model_diagram(document, appendix="D"):
    document.add_heading(f"Appendix {appendix}: Editable Data Model", level=1)
    document.add_paragraph(f"Figure {appendix}1. Proposed CareerFit data relationships. The model is editable in Word.")
    table = document.add_table(rows=7, cols=3)
    style_table(table)
    cells = [row.cells for row in table.rows]
    entities = (
        ("User", "1", "Authentication identity"),
        ("User Profile", "1:1 with User", "Preferences and autofill"),
        ("Resume", "many per User", "Editable resume versions"),
        ("Match Report", "many per User", "Comparison snapshots"),
        ("Job Description", "many per User", "Selected and tracked roles"),
        ("Search Alert", "many per User", "Reusable search filters"),
    )
    for index, (entity, relation, purpose) in enumerate(entities):
        set_cell_text(cells[index][0], entity, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[index][1], relation, color=TEAL, size=8)
        set_cell_text(cells[index][2], purpose, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(cells[index][0], LIGHT_BLUE if index % 2 == 0 else LIGHT_TEAL)
    cells[6][0].merge(cells[6][2])
    set_cell_text(
        cells[6][0],
        "Match Report links a resume snapshot and a job-description snapshot so historical results remain understandable after later edits.",
        size=8,
    )
    set_cell_shading(cells[6][0], LIGHT_AMBER)
    document.add_paragraph()


def add_wireframe(document, title, rows):
    document.add_heading(title, level=2)
    table = document.add_table(rows=len(rows), cols=3)
    style_table(table)
    for row_index, row in enumerate(rows):
        values = row if isinstance(row, tuple) else (row, "", "")
        if len(values) == 1:
            values = (values[0], "", "")
        for column_index, value in enumerate(values):
            set_cell_text(
                table.cell(row_index, column_index),
                value,
                bold=row_index == 0,
                size=8,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            set_cell_shading(table.cell(row_index, column_index), LIGHT_SLATE if row_index else LIGHT_TEAL)
    document.add_paragraph()


def add_mockup_notes(document, notes):
    table = document.add_table(rows=len(notes), cols=2)
    style_table(table)
    for row_index, (label, detail) in enumerate(notes):
        set_cell_text(table.cell(row_index, 0), label, bold=True, color=TEAL, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(row_index, 1), detail, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(table.cell(row_index, 0), LIGHT_TEAL)
    document.add_paragraph()


def add_wireframes(document, appendix="E"):
    document.add_heading(f"Appendix {appendix}: Editable Annotated Screen Mockups", level=1)
    document.add_paragraph(
        "These annotated desktop mockups describe the planned prototype before implementation. "
        "They are intentionally editable in Word so that the layout and labels can be revised after the midterm review."
    )
    add_wireframe(document, f"Figure {appendix}1. Home Page Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Log in | Sign up"),
        ("JOB SEARCH AND RESUME READINESS", "", "Hero image: resume workspace"),
        ("Turn each job posting into a stronger application.", "", "Report preview: 78%"),
        ("Upload your resume, find real roles, and see the most useful improvements before you apply.", "", "ATS structure: 8/9"),
        ("CHECK YOUR RESUME FIT", "SEARCH JOBS", "Skills coverage: 12/16"),
        ("1. Add resume", "2. Choose job", "3. Review fit | 4. Improve and track"),
        ("Private workspace", "Explainable scores", "Free to start"),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Introduce the workflow and provide immediate entry points."),
        ("Primary action", "Check your resume fit."),
        ("Design rationale", "The first viewport shows the product value, next action, and an explainable-report preview."),
    ])
    add_wireframe(document, f"Figure {appendix}2. Profile Page Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Signed-in user"),
        ("Candidate profile", "", "Profile completion"),
        ("Name", "Email", "Phone"),
        ("Location", "Target role", "Experience level"),
        ("Workplace preference", "Professional summary", ""),
        ("SAVE PROFILE", "Saved confirmation / validation message", ""),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Capture candidate preferences and reduce repeated search input."),
        ("Autofill", "Target role and location prefill job search fields while remaining editable."),
        ("Feedback", "Display save confirmation and field-specific validation messages."),
    ])
    add_wireframe(document, f"Figure {appendix}3. Resume Workspace Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Load sample"),
        ("Resume workspace: Add the resume you want to improve", "", ""),
        ("Upload PDF / DOCX / TXT", "Editable resume text preview", "Saved versions"),
        ("Uploading / success / error feedback", "Text editor with extracted content", "Version name | Save"),
        ("Accepted file types and 5 MB limit", "Word count | Character count", "Load version | Delete"),
        ("", "CLEAR RESUME | CONTINUE TO JOBS", "ATS preparation: 7/9"),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Create a reliable, editable resume source before comparison."),
        ("Recovery", "A wrong file can be dismissed, cleared, replaced, or corrected manually."),
        ("ATS preview", "Show document-preparation checks before the detailed job-specific report."),
    ])
    add_wireframe(document, f"Figure {appendix}4. Jobs Page Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Load sample"),
        ("Job discovery: Choose the role you want to evaluate", "", ""),
        ("Import public job URL", "URL input", "IMPORT URL"),
        ("Job title", "Location", "Country | SEARCH JOBS"),
        ("Filters: workplace | skills | exclude keywords | experience | employment | salary", "", "SAVE SEARCH ALERT"),
        ("Role insights: recurring skills | related role suggestions", "", ""),
        ("Search results and pagination", "Job cards: company | location | salary | source | freshness", "Selected job"),
        ("OPEN POSTING | SAVE ROLE | USE THIS JOB", "", "Quick match | readiness | skills"),
        ("COMPARE UP TO 3 JOBS", "Side-by-side role | company | salary | workplace | freshness", ""),
        ("", "", "GENERATE FULL REPORT"),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Search or import a role and make the comparison path obvious."),
        ("Without resume", "Selecting a job opens a prompt with a direct upload-resume action."),
        ("With resume", "Selecting a job shows a lightweight comparison immediately."),
    ])
    add_wireframe(document, f"Figure {appendix}5. Detailed Report Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Print | Save"),
        ("Readiness report: Selected role and company", "", "Add to tracker"),
        ("Match score", "Readiness score", "ATS preparation"),
        ("Executive summary", "Top priorities", "Score explanation"),
        ("Matched requirements", "Partial / weak evidence", "Missing requirements"),
        ("Requirement | priority | resume evidence", "", ""),
        ("ATS checklist", "Specific improvement examples", "Resume template draft"),
        ("Interview preparation", "Optional AI coaching", "RESCAN AFTER EDITS"),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Turn a score into an understandable revision plan."),
        ("Reading order", "Summary first, then prioritized improvements, then detailed evidence."),
        ("AI boundary", "Optional coaching is visually separated from deterministic results."),
    ])
    add_wireframe(document, f"Figure {appendix}6. Application Tracker Mockup", [
        ("CareerFit | Home | Profile | Resume | Jobs | Report | Tracker", "", "Export CSV"),
        ("Application tracker", "Search and filters", "Add role"),
        ("Saved", "Preparing", "Applied"),
        ("Interview", "Offer", "Rejected / Archived"),
        ("Selected application details", "Notes and tasks", "Dates and recruiter"),
        ("Linked resume version", "Cover letter and email drafts", "SAVE CHANGES"),
    ])
    add_mockup_notes(document, [
        ("Primary purpose", "Keep application progress and follow-up actions together."),
        ("Record detail", "Preserve notes, dates, recruiter contact, tasks, and linked resume version."),
        ("Progress model", "Use clear stages from saved role through archived outcome."),
    ])
    add_wireframe(document, f"Figure {appendix}7. Mobile Layout Mockup", [
        ("CareerFit", "Menu", "Account"),
        ("Page title and short guidance", "", ""),
        ("Primary action", "", ""),
        ("Main content stacked vertically", "", ""),
        ("Expandable secondary sections", "", ""),
        ("Sticky or visible next-step action", "", ""),
    ])
    add_mockup_notes(document, [
        ("Responsive rule", "Multi-column desktop layouts become a single reading column."),
        ("Navigation", "Horizontal navigation becomes a compact menu."),
        ("Priority", "The current task and next action remain visible before secondary detail."),
    ])


def configure_document(document):
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    for style_name, size, color in (
        ("Heading 1", 15, TEAL),
        ("Heading 2", 12, INK),
        ("Heading 3", 11, INK),
    ):
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        add_page_number(section.footer.paragraphs[0])


def main():
    markdown = SOURCE_PATH.read_text(encoding="utf-8")
    markdown = markdown[markdown.index("## 1. Introduction"):]
    document = Document()
    configure_document(document)
    add_cover_page(document)

    document.add_heading("Table of Contents", level=1)
    add_toc(document.add_paragraph())
    document.add_page_break()

    add_markdown_content(document, markdown)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_architecture_diagram(document)
    add_workflow_diagram(document)
    add_navigation_diagram(document)
    add_data_model_diagram(document)
    add_wireframes(document)
    document.save(OUTPUT_PATH)
    print(f"Generated {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
