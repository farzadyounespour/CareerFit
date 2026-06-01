from io import BytesIO
import re

from docx import Document
from pypdf import PdfReader


class ResumeParseError(ValueError):
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_resume_text(uploaded_file):
    filename = uploaded_file.name or ""
    extension = _get_extension(filename)

    if extension == ".pdf":
        text = _extract_pdf_text(uploaded_file)
    elif extension == ".docx":
        text = _extract_docx_text(uploaded_file)
    elif extension == ".txt":
        text = _extract_plain_text(uploaded_file)
    else:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ResumeParseError(f"Unsupported resume file type. Use {supported}.")

    normalized = _normalize_extracted_text(text)
    if not normalized:
        raise ResumeParseError("No readable text was found in this resume file.")

    return normalized


def _get_extension(filename):
    lowered = filename.lower()
    for extension in SUPPORTED_EXTENSIONS:
        if lowered.endswith(extension):
            return extension
    return ""


def _extract_pdf_text(uploaded_file):
    try:
        reader = PdfReader(BytesIO(uploaded_file.read()))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ResumeParseError("Unable to read this PDF resume.") from exc


def _extract_docx_text(uploaded_file):
    try:
        document = Document(BytesIO(uploaded_file.read()))
    except Exception as exc:
        raise ResumeParseError("Unable to read this DOCX resume.") from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def _extract_plain_text(uploaded_file):
    raw = uploaded_file.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParseError("Unable to decode this text resume.")


def _normalize_extracted_text(text):
    lines = [_normalize_extracted_line(line) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _normalize_extracted_line(line):
    fragments = re.split(r"\s{2,}", line.strip())
    return " ".join(
        repaired
        for fragment in fragments
        if (repaired := _repair_spaced_fragment(fragment))
    )


def _repair_spaced_fragment(fragment):
    tokens = fragment.split()
    if len(tokens) < 2:
        return " ".join(tokens)

    single_character_tokens = sum(
        len(token.strip("(),.%+/")) <= 1
        for token in tokens
    )
    if single_character_tokens / len(tokens) < 0.7:
        return " ".join(tokens)

    return re.sub(r"\s+", "", fragment)
